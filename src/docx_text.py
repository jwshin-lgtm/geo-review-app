"""docx 바이트에서 문단 텍스트를 뽑아내는 유틸."""
from __future__ import annotations

import io

from docx import Document


def load_document(docx_bytes: bytes) -> Document:
    return Document(io.BytesIO(docx_bytes))


def extract_paragraphs(docx_bytes: bytes, include_tables: bool = False) -> list[str]:
    """본문 문단 텍스트 목록. 빈 문단도 순서 유지를 위해 포함한다."""
    doc = load_document(docx_bytes)
    paragraphs = [p.text for p in doc.paragraphs]

    if include_tables:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(p.text for p in cell.paragraphs)

    return paragraphs


def extract_nonempty_paragraphs(docx_bytes: bytes) -> list[str]:
    return [text for text in extract_paragraphs(docx_bytes) if text.strip()]
