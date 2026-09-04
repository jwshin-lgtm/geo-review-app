"""원본 docx를 열어, 바뀐 문단의 run만 교체하고 변경분에 하이라이트를 입힌다."""
from __future__ import annotations

import io

from docx.enum.text import WD_COLOR_INDEX

from . import diff_utils
from .docx_text import load_document

RunFormat = dict


def get_run_format(run) -> RunFormat:
    if run is None:
        return {}
    color = run.font.color
    return {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "font_name": run.font.name,
        "font_size": run.font.size,
        "color_rgb": color.rgb if color and color.type is not None else None,
    }


def apply_run_format(run, fmt: RunFormat) -> None:
    run.bold = fmt.get("bold")
    run.italic = fmt.get("italic")
    run.underline = fmt.get("underline")
    if fmt.get("font_name"):
        run.font.name = fmt["font_name"]
    if fmt.get("font_size"):
        run.font.size = fmt["font_size"]
    if fmt.get("color_rgb") is not None:
        run.font.color.rgb = fmt["color_rgb"]


def build_run_map(paragraph) -> list[tuple[int, int, object]]:
    """[(start_offset, end_offset, run), ...] - run.text 길이 누적 기준."""
    run_map = []
    offset = 0
    for run in paragraph.runs:
        length = len(run.text)
        run_map.append((offset, offset + length, run))
        offset += length
    return run_map


def _run_at(run_map: list[tuple[int, int, object]], pos: int):
    for start, end, run in run_map:
        if start <= pos < end:
            return start, end, run
    if run_map:
        return run_map[-1]
    return None


def slice_by_original_runs(text_segment: str, start_pos: int, run_map: list[tuple[int, int, object]]):
    """text_segment(원문 그대로의 구간)를, 걸쳐있는 원본 run 경계에 맞춰 (부분텍스트, run) 리스트로 쪼갠다."""
    segments = []
    pos = start_pos
    idx = 0
    n = len(text_segment)
    while idx < n:
        found = _run_at(run_map, pos)
        if found is None:
            segments.append((text_segment[idx:], None))
            break
        run_start, run_end, run = found
        take = min(run_end - pos, n - idx)
        take = max(take, 1)
        segments.append((text_segment[idx : idx + take], run))
        idx += take
        pos += take
    return segments


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def apply_highlighted_revision(doc, para_index: int, revised_text: str) -> bool:
    """paragraph를 수정된 텍스트로 교체하고, 바뀐/추가된 부분만 하이라이트한다.

    반환값: 실제로 내용이 바뀌었는지 여부.
    """
    paragraph = doc.paragraphs[para_index]
    original_text = paragraph.text
    if original_text == revised_text:
        return False

    run_map = build_run_map(paragraph)
    base_run = paragraph.runs[0] if paragraph.runs else None
    base_fmt = get_run_format(base_run)

    old_tokens, new_tokens, opcodes = diff_utils.word_diff_opcodes(original_text, revised_text)

    clear_runs(paragraph)

    orig_pos = 0
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            segment_text = "".join(old_tokens[i1:i2])
            for sub_text, run in slice_by_original_runs(segment_text, orig_pos, run_map):
                new_run = paragraph.add_run(sub_text)
                apply_run_format(new_run, get_run_format(run) if run else base_fmt)
            orig_pos += len(segment_text)
        else:
            if tag in ("replace", "insert"):
                segment_text = "".join(new_tokens[j1:j2])
                if segment_text:
                    new_run = paragraph.add_run(segment_text)
                    apply_run_format(new_run, base_fmt)
                    new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if tag in ("replace", "delete"):
                orig_pos += len("".join(old_tokens[i1:i2]))

    return True


def build_highlighted_docx(original_docx_bytes: bytes, revisions: list[dict]) -> tuple[bytes, bool]:
    """revisions: reviser.revise_paragraphs()의 결과. 반환: (수정된 docx bytes, 실제 변경 여부)."""
    doc = load_document(original_docx_bytes)
    changed_any = False
    for rev in revisions:
        if rev.get("changed"):
            changed = apply_highlighted_revision(doc, rev["index"], rev["revised_text"])
            changed_any = changed_any or changed

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), changed_any
